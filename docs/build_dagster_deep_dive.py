from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).with_name('Dagster_Orchestration_Deep_Dive_Guide.docx')

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(9)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, 'FFFFFF'); shade(t.rows[0].cells[i], '1F4D78')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],v)
    if widths:
        for row in t.rows:
            for cell,w in zip(row.cells,widths): cell.width=Inches(w)
    doc.add_paragraph()
    return t
def bullet(doc, text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(text); return p
def heading(doc, text, level=1): doc.add_heading(text, level=level)
def note(doc, title, text):
    t=doc.add_table(rows=1, cols=1); t.style='Table Grid'; c=t.cell(0,0); shade(c,'E8EEF5'); c.text=''
    p=c.paragraphs[0]; r=p.add_run(title+': '); r.bold=True; r.font.color.rgb=RGBColor(31,77,120); p.add_run(text)
    doc.add_paragraph()

doc=Document(); sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Inches(0.75); sec.left_margin=sec.right_margin=Inches(0.8)
styles=doc.styles; styles['Normal'].font.name='Calibri'; styles['Normal'].font.size=Pt(10); styles['Normal'].paragraph_format.space_after=Pt(5); styles['Normal'].paragraph_format.line_spacing=1.12
for name,size,color in [('Title',20,'1F4D78'),('Heading 1',15,'2E74B5'),('Heading 2',12,'1F4D78'),('Heading 3',11,'1F4D78')]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Dagster Orchestration Deep-Dive Guide'); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=RGBColor(31,77,120)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Credit Card Fraud Detection Data Engineering Pipeline').italic=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('AWS S3 → Databricks Bronze/Silver → Snowflake RAW/Gold → Streamlit').font.color.rgb=RGBColor(90,90,90)
doc.add_paragraph()
note(doc,'Scope','This guide describes the implemented Dagster project, not a generic design. It explains how the local Dagster daemon, the two sensors, S3, Databricks Jobs, audit tables, Snowflake, and file routing work together.')

heading(doc,'1. What Dagster is doing in this project')
doc.add_paragraph('Dagster is the control plane. It does not transform the CSV itself. Instead, it detects a file, creates one orchestration run, passes parameters to Databricks Jobs, waits for each stage, validates the Snowflake load, refreshes Gold, and then archives the source file.')
table(doc,['Execution order','What happens','Primary evidence'],[
['1. Landing sensor','Polls S3 landing/ and chooses the oldest unseen file.','Dagster sensor tick, cursor and RunRequest'],
['2. Bronze job','Databricks reads the selected S3 file, preserves raw rows and writes file audit data.','workspace.bronze.transactions_raw; workspace.ops.processed_files'],
['3. Silver job','Databricks applies the approved standardisation/quality transformations for that batch.','workspace.silver.transactions_silver_incremental; silver audit'],
['4. Snowflake RAW load','The selected Silver batch is appended only if the same batch/source is not already present.','FRAUD_DB.RAW.TRANSACTIONS'],
['5. RAW validation','Dagster queries Snowflake to ensure the target batch has rows.','Snowflake count for BATCH_ID'],
['6. Gold refresh','Dagster calls ANALYTICS.REFRESH_FRAUD_GOLD().','Gold tables/views and Streamlit data'],
['7. Archive','Only after Gold succeeds, copy landing object to archive/ then delete landing object.','S3 archive object and Dagster materialization'],
],[1.25,3.75,2.0])
note(doc,'Key rule','One Dagster pipeline run processes one source file and one batch. Its op dependencies force the stages to run in sequence: Bronze → Silver → Snowflake RAW → validation → Gold → archive.')

heading(doc,'2. How a Dagster sensor works internally')
doc.add_paragraph('A sensor is a small function evaluated periodically by the Dagster daemon. Every evaluation is called a tick. The sensor can either return a SkipReason (nothing should start) or a RunRequest (start the job with configuration). It is not a continuously open S3 connection and it does not receive an AWS push event in this project; it polls S3.')
table(doc,['Step inside a tick','Current implementation'],[
['1. Daemon invokes sensor','landing_file_sensor has a minimum interval of 30 seconds. Dagster evaluates it at approximately that frequency.'],
['2. Read active runs','It checks the Dagster instance for QUEUED, STARTING or STARTED fraud_pipeline_job runs. If one exists, it skips; this prevents two pipelines running simultaneously.'],
['3. List S3 objects','boto3 list_objects_v2 paginator lists objects under S3_LANDING_PREFIX (normally landing/). Folder placeholders are ignored.'],
['4. Compare with cursor','For each object it builds key:etag. If that key is in the saved sensor cursor, it is already seen.'],
['5. Choose one object','It chooses the oldest unseen candidate, not every candidate at once. Other files remain for later ticks.'],
['6. Apply duplicate guard','It first looks up a known ETag in Databricks audit; if unknown, it calculates SHA-256 and checks the audit table for successful identical content.'],
['7. Request or skip run','A unique file returns RunRequest with Bronze parameters. Duplicate content is copied to reject/duplicates/, audit-recorded, then skipped.'],
['8. Save cursor','The sensor saves up to the latest 500 key:etag values in Dagster instance storage.'],
],[1.65,5.35])

heading(doc,'3. The two sensors and why both are needed')
table(doc,['Sensor','Trigger and frequency','Responsibility','What it does not do'],[
['landing_file_sensor','Polling sensor; minimum interval 30 sec; default STOPPED until enabled.','Detects one new file in landing/, prevents concurrent pipeline starts, checks duplicate content, then launches fraud_pipeline_job.','It does not transform data and it does not archive a file before all stages succeed.'],
['pipeline_failure_reject_sensor','Run-status sensor; reacts to final FAILURE; min interval 15 sec; default STOPPED until enabled.','After all job retry attempts fail, moves the source file from landing/ to reject/validation_failed/.','It does not execute on a temporary retry and it does not repair the data automatically.'],
],[1.55,1.7,2.55,1.25])
note(doc,'Why failure routing waits','Every Dagster operation has RetryPolicy(max_retries=2, delay=30). The failure sensor runs only when the entire pipeline has finally failed. This avoids rejecting a file while Dagster is still correctly retrying a temporary Databricks/Snowflake outage.')

heading(doc,'4. How Dagster knows a file is new')
heading(doc,'4.1 Sensor cursor: short-term discovery memory',2)
doc.add_paragraph('The cursor stores a comma-separated list of key:etag values in Dagster’s own instance storage, not in S3 and not in Databricks. Example: landing/creditcard_batch_004.csv:abc123.... A repeated tick sees this same pair and skips it. The code keeps only the newest 500 entries, so the cursor is a convenience and concurrency guard—not the durable duplicate-control system.')
bullet(doc,'If the sensor is reset, its cursor becomes empty. It may see old objects again, but the durable audit/SHA check still blocks already-successful content.')
bullet(doc,'S3_SENSOR_START_AFTER can be set when enabling a sensor for the first time. Older objects are added to the cursor as seen without creating runs. This prevents historical landing files from starting unexpectedly.')
heading(doc,'4.2 ETag: S3 object-version fingerprint',2)
doc.add_paragraph('S3 returns an ETag for each object in ListObjectsV2. The sensor removes the quote marks and uses it for the run key and as the base of the batch ID. For a simple one-part upload it is often derived from the object MD5, but it must not be treated as a guaranteed content MD5: multipart uploads and some encryption/upload methods can produce different ETags for identical bytes.')
table(doc,['ETag use','Justification'],[
['Sensor run key: key:etag','Detects a changed/re-uploaded object at a location and makes Dagster’s RunRequest idempotent for that object version.'],
['Fast audit lookup','If a successful audit row has the same source_etag, the stored SHA-256 can be reused; the sensor avoids downloading a large file just to hash it again.'],
['Batch ID basis','Current code creates s3_<first 12 ETag characters>. This is a compact trace label and is passed to every downstream stage.'],
],[2.35,4.7])
heading(doc,'4.3 SHA-256: canonical byte-for-byte content identity',2)
doc.add_paragraph('SHA-256 is the decisive duplicate check. When the ETag has not been seen before, the sensor calls S3 GetObject, creates hashlib.sha256(), reads the object body in 1 MB chunks, updates the digest for each chunk, closes the body, and returns a 64-character hexadecimal hash. The full object is not loaded into memory at once.')
table(doc,['Question','Answer'],[
['When is it calculated?','Only for a new/unknown ETag. For a known successful ETag, the sensor retrieves the existing content_sha256 from Databricks audit as a fast path.'],
['Where is it stored?','Durably in workspace.ops.processed_files.content_sha256. It is recorded for successful Bronze ingestion and for duplicate content routed to reject/duplicates/. It is also passed as a Databricks notebook parameter.'],
['What exact comparison is made?','The sensor queries processed_files for the same content_sha256 with status = SUCCESS and target_table = workspace.bronze.transactions_raw.'],
['Why SHA-256 rather than filename?','A filename can be changed. SHA-256 represents the exact file bytes, so copied/renamed identical content still matches.'],
['Important limitation','A file with the same business rows but changed row order, line endings, header formatting, or extra whitespace has different bytes and a different SHA. A future enhancement would be a canonical row-level content fingerprint.'],
],[2.2,4.85])

heading(doc,'5. Duplicate prevention: defense in depth')
table(doc,['Layer','Control','What it prevents'],[
['Dagster active-run check','Sensor skips when a pipeline run is already QUEUED/STARTING/STARTED.','Two files starting two pipelines concurrently.'],
['Sensor cursor','Saved key:etag entries prevent repeated discovery on later ticks.','Same S3 object version repeatedly launching.'],
['ETag lookup','Known ETag maps to prior successful SHA.','Unnecessary re-hashing/reprocessing of unchanged object version.'],
['SHA-256 audit check','Exact file hash is checked against SUCCESS Bronze audit records.','Same data uploaded under a different filename or ETag.'],
['Databricks Bronze audit','processed_files records the selected file and outcome.','Untraceable/ambiguous file history.'],
['Snowflake RAW guard','Load checks exact BATCH_ID + SOURCE_FILE before append, then reconciles the count.','A pipeline retry silently appending the same batch twice.'],
['Archive/reject movement','Successful file moves to archive; duplicate/failed files move out of landing.','The same landing object being offered indefinitely.'],
],[1.45,2.65,2.95])
note(doc,'Important distinction','File-level duplicate prevention happens before Bronze. Bronze should preserve all source rows. Do not use SELECT DISTINCT in Bronze: it caused the earlier 1,081-row discrepancy. If business-approved row-level duplicate cleansing is required, perform and document it in Silver, not in the raw Bronze preservation layer.')

heading(doc,'6. Batch ID, pipeline identity and audit correlation')
doc.add_paragraph('The current sensor creates batch_id as s3_<first 12 characters of ETag>. Example: ETag 510c71c5655bc... becomes s3_510c71c5655b. The same batch ID is passed into Bronze, Silver and Snowflake load notebooks. It is a correlation key—not a claim that ETag is a cryptographic identifier. SHA-256 remains the content identity.')
table(doc,['Identifier','How it is generated','Where used','Why useful'],[
['batch_id','Sensor: s3_ + source ETag first 12 chars.','Bronze/Silver rows, file audits, Snowflake RAW, Gold validation, archive materialization.','Groups all processing evidence for one input object.'],
['pipeline_run_id','Bronze notebook: typically bronze_<batch_id> (or an execution-specific run value where configured).','Bronze rows and processed_files audit.','Correlates the table write with the ingestion attempt.'],
['Dagster run ID','Created by Dagster for every orchestration launch.','Databricks Job idempotency tokens: <dagster-run-id>-bronze/silver/snowflake-raw.','Prevents duplicate Databricks Jobs for retries within one Dagster orchestration run.'],
['Databricks run ID','Returned from Jobs API after run-now.','Dagster logs and asset materializations.','Links a Dagster stage to the actual Databricks workload.'],
],[1.35,2.25,2.1,1.35])

heading(doc,'7. Audit-column dictionary')
heading(doc,'7.1 Row-level provenance carried in Bronze (and into Silver)',2)
table(doc,['Column','Created','Meaning and use'],[
['load_ts','current_timestamp() while Bronze batch view is created.','Timestamp when this pipeline loaded the record into Bronze. Used for lineage and operational investigation; it is not the transaction business time.'],
['source_file','Sensor passes source_file_name; Bronze writes it as a literal on every row.','Original file name, useful to filter/reconcile a file batch.'],
['source_s3_uri','Sensor constructs s3://bucket/key and passes source_path.','Exact S3 source address, useful for trace-back and archive/reject investigation.'],
['batch_id','Sensor creates s3_<ETag prefix>.','Batch correlation key across Bronze, Silver, Snowflake RAW and Dagster evidence.'],
['ingestion_date','current_date() during Bronze load.','Partition/reporting-friendly date the pipeline ingested the record; not the transaction date.'],
['pipeline_run_id','Bronze creates a pipeline-specific identifier, e.g. bronze_<batch_id>.','Correlates each record with its ingestion execution/audit event.'],
],[1.45,2.35,3.9])
heading(doc,'7.2 File-level audit: workspace.ops.processed_files',2)
table(doc,['Column','Creation / meaning','Why it matters'],[
['source_file_name','Final part of S3 key passed by sensor.','Human-readable file trace.'],
['source_s3_uri / source_path','Sensor builds s3://bucket/key.','Exact source object used for a run.'],
['source_etag','S3 ListObjectsV2 item ETag, quote marks stripped.','Fast object-version lookup; not the final content proof.'],
['content_sha256','Sensor streams object bytes through hashlib.sha256.','Durable byte-level duplicate prevention across renamed files.'],
['file_size_bytes','S3 object Size returned in ListObjectsV2.','Quick file-volume evidence and sanity check.'],
['batch_id','Sensor-generated ETag-based correlation ID.','Ties audit row to downstream batch data.'],
['pipeline_run_id','Bronze execution identifier.','Trace an attempt even when a file is rerun/troubleshot.'],
['status','Lifecycle state: RUNNING, SUCCESS, FAILED; duplicate route is recorded as DUPLICATE_CONTENT.','Determines whether content can block future loads and communicates outcome. Only SUCCESS is used as an already-processed content match.'],
['source_row_count','Count read from source before Bronze write.','Input reconciliation.'],
['bronze_row_count','Count written for the Bronze batch.','Proves Bronze preservation/reconciliation.'],
['started_ts / completed_ts','Timestamps written around the attempt.','Duration, ordering, SLA evidence and stuck RUNNING diagnosis.'],
['target_table','Fully qualified intended table, e.g. workspace.bronze.transactions_raw.','Avoids treating a successful load to a different target as the same event.'],
['error_message','Captured exception/message on failure.','First place to diagnose a stage failure.'],
['rejection_reason / routed_s3_uri','Set for duplicate content routing.','Explains why a file was rejected and where it was moved.'],
],[1.55,3.1,3.05])
heading(doc,'7.3 Silver and Snowflake audit evidence',2)
table(doc,['Evidence','What it contains','Use'],[
['workspace.ops.silver_processed_batches','source file, batch, status, Bronze count, Silver count, started/completed timestamps, error message.','Shows that the expected Bronze batch became the expected Silver batch.'],
['FRAUD_DB.RAW.TRANSACTIONS','Silver business columns plus carried provenance such as BATCH_ID and SOURCE_FILE.','Supports exact per-batch Snowflake reconciliation and Gold refresh.'],
['Dagster run logs/materializations','Dagster run ID, Databricks Job IDs, asset events and stage state.','Operational proof of the whole orchestration chain.'],
],[1.75,3.6,2.35])

heading(doc,'8. What happens when a stage fails')
doc.add_paragraph('A failure does not automatically mean duplicate data. The correct answer depends on the stage and whether the write committed. The safe rule is: never blindly rerun a failed file. First inspect the stage audit, target row count and error message.')
table(doc,['Failure point','Automatic behavior','Operator recovery / duplication safety'],[
['Sensor / pre-Bronze','No pipeline starts; tick may report an error.','Fix AWS credentials, Databricks SQL Warehouse access or missing audit table; then re-enable/retry sensor. No table rows have been written.'],
['Bronze job','Dagster retries up to 2 times with 30-sec delay. Final failure triggers validation_failed reject routing.','Check processed_files and Bronze count for the batch. If write did not commit, retry is safe. If table rows exist but audit says RUNNING/FAILED, reconcile first; do not blindly append.'],
['Silver job','Bronze already exists; Dagster retries Silver.','Check silver_processed_batches and Silver batch count. Use batch_id/source_file to identify whether the batch was completed; fix transformation/schema issue then rerun only after confirming the target state.'],
['Snowflake RAW load','Guard checks BATCH_ID + SOURCE_FILE before append and reconciles row count.','If 0 rows, retry is safe. If partial rows exist, do not blindly append: verify count and remove/reload only that exact batch+source after evidence is captured.'],
['Snowflake validation','Gold is blocked when count is zero.','Resolve RAW issue first. Gold is not refreshed until validation passes.'],
['Gold refresh','Dagster retries procedure.','Gold uses CREATE OR REPLACE aggregate tables/views, so it is idempotent: rerunning refresh recalculates output rather than appending duplicate metrics.'],
['Archive','Data stages succeeded; copy occurs before delete.','Retry archive only. Copying same object to same archive key is safe; a missing landing source after a prior move is handled as already moved.'],
],[1.45,2.4,3.85])
note(doc,'Recovery priority','1) Pause/disable the landing sensor if a bad file keeps being offered. 2) Find its audit row by source file/path/batch. 3) Compare row counts in the target layer. 4) Fix the root cause. 5) Resume only the safe stage or rerun the full job if no write committed. 6) Verify audit status and file location afterward.')

heading(doc,'9. Practical scenarios and defensible answers')
table(doc,['Scenario','Expected outcome','Why'],[
['Same filename uploaded again, unchanged object/ETag','Cursor or ETag audit fast path stops it; no new data load.','It is already seen and/or maps to a successful SHA.'],
['Different filename, exact same bytes','SHA-256 matches successful audit, object moves to reject/duplicates/, status is DUPLICATE_CONTENT.','Filename is not trusted as identity; exact content is.'],
['Same business data but bytes differ','It may be treated as new because SHA hashes bytes.','Row order, line endings and header formatting change the hash. Canonical content fingerprint is a future enhancement.'],
['Sensor restarted / cursor reset','Object may be discovered again, but audit SHA guard still blocks success content.','Cursor is not the only protection.'],
['File fails after Bronze','Dagster retries; after final pipeline failure, failure sensor routes file to reject/validation_failed/.','Failure routing happens only after retry exhaustion.'],
['Multiple new files land together','Only oldest unseen file starts; active-run check blocks another pipeline until it finishes.','Avoids competing writes and makes demo lineage easier to follow.'],
],[2.2,2.6,2.9])

heading(doc,'10. Demo and troubleshooting checklist')
table(doc,['Check','Where / example'],[
['Sensor is running','Dagster Automation → landing_file_sensor; check latest tick and cursor.'],
['No concurrent pipeline','Dagster Runs; only one active fraud_pipeline_job should be present.'],
['Bronze source reconciliation','Query processed_files with selected batch_id + source_file_name; compare source_row_count and bronze_row_count.'],
['Silver reconciliation','Query workspace.ops.silver_processed_batches by batch_id.'],
['Snowflake RAW reconciliation','SELECT COUNT(*) FROM FRAUD_DB.RAW.TRANSACTIONS WHERE BATCH_ID = ''<batch_id>'';'],
['Gold refreshed','CALL ANALYTICS.REFRESH_FRAUD_GOLD(); or show Dagster refresh_gold event and Gold table timestamps.'],
['File route','S3: archive/<file> for success; reject/duplicates/<file> for same content; reject/validation_failed/<file> for final failure.'],
['Diagnose sensor 403','Verify DATABRICKS_TOKEN scope/user, DATABRICKS_SQL_WAREHOUSE_ID, warehouse permission, and that workspace.ops.processed_files exists.'],
],[2.3,5.4])

heading(doc,'11. Security and operating precautions')
bullet(doc,'Keep AWS keys, Databricks personal access tokens and Snowflake password only in .env / secret storage. Never commit them to Git.'); bullet(doc,'The Dagster AWS IAM user needs least-privilege list/get for landing and put/delete only for archive and reject paths.'); bullet(doc,'The Databricks token is used both for Jobs API and Databricks SQL Statement Execution API. SQL Warehouse access is required for sensor audit checks.'); bullet(doc,'Treat processed_files as an operational control table. Do not delete it while the sensor is enabled; recreate Bronze audit DDL before enabling the sensor after any reset.'); bullet(doc,'Never use the historical Class label to create a suspicious-risk outcome. It is retained only for historical validation/reporting where approved.');

heading(doc,'12. One-minute explanation for review')
doc.add_paragraph('“Dagster polls S3 landing every 30 seconds. It saves key-and-ETag values in its cursor, processes only one oldest unseen file at a time, and prevents simultaneous pipeline runs. Before Bronze, it checks the object’s ETag against the Databricks audit; for a new ETag it streams the file through SHA-256. If the exact content was successfully processed previously, it routes the file to reject/duplicates and records the reason. Otherwise it passes the file details and batch ID into the Bronze job. Bronze and Silver retain provenance, Snowflake RAW is guarded by batch/source reconciliation, Gold is refreshed only after RAW validation, and only then is the input copied to archive and removed from landing. If the full pipeline fails after retries, a second run-status sensor moves the source to reject/validation_failed.”')

doc.add_paragraph('Document version: Final implementation guide • Generated 11 August 2026').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.save(OUT)
print(OUT)
