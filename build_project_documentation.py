from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('outputs/Credit_Card_Fraud_Detection_Project_Documentation.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = '0B2545'
BLUE = '2E74B5'
TEAL = '0F766E'
LIGHT_BLUE = 'E8EEF5'
LIGHT_TEAL = 'E6F4F1'
LIGHT_GOLD = 'FFF7E0'
GRAY = '5B6573'
WHITE = 'FFFFFF'
RED = '9B1C1C'
GREEN = '146C43'

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 16, 8),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 11.5, NAVY, 8, 4),
]:
    s = styles[name]
    s.font.name = 'Calibri'
    s._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    s._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def table_borders(table, color='C7D1DE'):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def write_cell(cell, text, bold=False, color='000000', size=9.3):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = 'Calibri'
    r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)

def add_table(headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table_borders(table)
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, header_fill)
        write_cell(cell, text, bold=True, color=NAVY, size=9.2)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            write_cell(cell, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_title(text, size=29):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.2*level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p

def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p

def add_callout(label, text, fill=LIGHT_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table_borders(table, 'B9D8D3')
    set_table_geometry(table, [9120])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(label + ': ')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_code_flow(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    shade = OxmlElement('w:shd')
    shade.set(qn('w:fill'), 'F3F6FA')
    p._p.get_or_add_pPr().append(shade)

def footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run('Credit Card Fraud Detection | Capstone Documentation | August 2026')
    r.font.name = 'Calibri'; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GRAY)

footer(sec)

# Cover
add_title('Credit Card Fraud Detection')
add_subtitle('End-to-End Data Engineering Pipeline — Project Documentation')
add_code_flow('AWS S3  ->  Databricks Bronze  ->  Databricks Silver  ->  Snowflake RAW  ->  Snowflake Gold  ->  Streamlit')
meta = add_table(
    ['Document item', 'Value'],
    [
        ('Project type', 'Data engineering capstone; rule-based fraud-review prioritisation'),
        ('Core tools', 'AWS S3, Databricks, Snowflake, Dagster, SQL, Python, Streamlit, Git'),
        ('Primary data set', 'Kaggle ULB Credit Card Fraud Detection (Time, V1–V28, Amount, Class)'),
        ('Submission deadline', '12 August 2026'),
        ('Document purpose', 'Explain the full implementation, operations, controls, and final outputs.'),
    ], [2200, 6920])
doc.add_paragraph()
add_callout('Important scope note', 'This is a data-engineering and operational review-prioritisation pipeline, not a machine-learning fraud model. Class/FRAUD_LABEL is retained only for historical validation and reporting; it is never used to create the review score or review priority.', LIGHT_GOLD)
doc.add_page_break()

doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph('The project builds an end-to-end, auditable data pipeline for the ULB credit-card transaction data set. Source files arrive in Amazon S3, are ingested into Databricks Bronze and transformed into a clean, analyst-ready Silver table. The Silver data is incrementally loaded to Snowflake RAW. Snowflake Gold tables and an analyst-review view support a native Streamlit dashboard. Dagster orchestrates the complete workflow, monitors new landing files, controls duplicate processing, validates loading results, and archives successfully completed files.')
add_callout('Business outcome', 'The pipeline does not assert that a transaction is fraud. It produces a defensible review priority (High, Medium, or Low) using observable transaction patterns so an analyst can examine the most relevant items first.')

doc.add_heading('2. Objective, Scope, and Data Limitations', level=1)
doc.add_heading('2.1 Objective', level=2)
doc.add_paragraph('Build a reliable and demonstrable data platform that ingests credit-card transactions, preserves raw data, derives review-oriented features, publishes curated analytics, and exposes the results to business users through a dashboard.')
doc.add_heading('2.2 What is being built', level=2)
for item in [
    'Incremental file ingestion from S3 landing into Databricks Bronze.',
    'Audited and idempotent processing using source metadata, file content SHA-256, batch identifiers, and process-status tables.',
    'Silver standardisation, exact-row deduplication, and transparent rule-based review features.',
    'Snowflake RAW storage, Gold analytics outputs, and an analyst review queue.',
    'Dagster orchestration with sequential dependencies, retries, validation, archive routing, and a new-file sensor.',
    'A Snowflake-native Streamlit dashboard for operational monitoring and review prioritisation.',
]: add_bullet(item)
doc.add_heading('2.3 Source-data limitations and design response', level=2)
add_table(['Source-data fact', 'Impact', 'Project response'], [
    ('V1–V28 are anonymised PCA features', 'They cannot be explained as customer, card, merchant, or location attributes.', 'They are preserved. A controlled repeated-anonymised-pattern signal may be used as a technical pattern indicator, not business identity logic.'),
    ('Time is elapsed seconds', 'It is not a calendar timestamp.', 'A clearly labelled synthetic transaction timestamp is derived from a defined project start date.'),
    ('No customer/card/merchant identifiers', 'Customer velocity, merchant trends, and genuine dimensions cannot be built.', 'Use time-window/repetition signals only; do not claim entity-level fraud logic.'),
    ('Class is the historical label', 'Using it to create a score would leak the answer.', 'Use it only to compare historical outcomes in Gold analytics and validate the design.'),
], [2360, 2700, 4060])

doc.add_heading('3. Final Architecture and Data Flow', level=1)
doc.add_paragraph('The final architecture has one required data path and a mandatory orchestration path. The full editable architecture figure is maintained in the repository at diagrams/final_fraud_detection_architecture_v2.svg.')
add_code_flow('Kaggle CSV -> S3 landing/ -> Bronze transactions_raw -> Silver transactions_silver_incremental -> Snowflake RAW.TRANSACTIONS -> Gold tables/view -> Snowflake Streamlit dashboard')
add_code_flow('Dagster sensor -> Bronze job -> Silver job -> Snowflake load job -> RAW validation -> Gold refresh -> archive/ (or reject/ on duplicate/failure)')
add_table(['Layer / component', 'Purpose', 'Primary final object(s)'], [
    ('AWS S3', 'File landing and lifecycle storage.', 'landing/, archive/, reject/duplicates/, reject/validation_failed/'),
    ('Databricks Bronze', 'Immutable-style raw preservation plus ingestion lineage.', 'workspace.bronze.transactions_raw'),
    ('Databricks Operations', 'File and batch audit records.', 'workspace.ops.processed_files; workspace.ops.silver_processed_batches'),
    ('Databricks Silver', 'Clean, exact-deduplicated, feature-enriched transactions.', 'workspace.silver.transactions_silver_incremental'),
    ('Snowflake RAW', 'Incremental analyst-ready transaction storage.', 'FRAUD_DB.RAW.TRANSACTIONS'),
    ('Snowflake Gold', 'Dashboard aggregates and real-time review presentation.', 'FRAUD_DAILY_SUMMARY; AMOUNT_BUCKET_SUMMARY; RISK_TIME_DASHBOARD; ANALYST_REVIEW_QUEUE'),
    ('Snowflake Streamlit', 'Interactive consumer application.', 'FRAUD_DB.ANALYTICS.FRAUD_RISK_DASHBOARD'),
], [2000, 3700, 3420])

doc.add_heading('4. Tooling and Responsibilities', level=1)
add_table(['Tool', 'How it is used', 'Control / reason'], [
    ('Amazon S3', 'Stores inbound, archived, and rejected CSV files.', 'Separates lifecycle states and supports secure programmatic access.'),
    ('Databricks + Unity Catalog', 'Runs SQL notebooks, stores Bronze/Silver Delta tables, and exposes governed catalog objects.', 'Catalog/schema naming makes data locations explicit.'),
    ('Snowflake', 'Receives incremental RAW data; builds Gold analytics; hosts Streamlit.', 'Separates RAW from business-consumption outputs.'),
    ('Dagster', 'Schedules and observes the end-to-end workflow.', 'Dependencies prevent downstream work starting before upstream success.'),
    ('Git/GitHub', 'Stores source code, notebook exports, SQL, diagrams, and documentation.', 'Provides version control and collaboration history.'),
    ('Python', 'Used for Dagster and the Snowflake connector workload.', 'Not used for the final Bronze/Silver transformations, which are SQL-first.'),
], [1700, 4100, 3320])

doc.add_heading('5. S3 Landing and File Lifecycle', level=1)
doc.add_paragraph('The bucket intern-final-project-fraud-detection is in ap-south-1. S3 uses prefixes rather than physical folders; the following prefixes make the operational file state visible.')
add_table(['S3 prefix', 'Meaning', 'What happens there'], [
    ('landing/', 'New source files awaiting processing.', 'Dagster sensor detects eligible CSV files and starts one pipeline run.'),
    ('archive/', 'Successfully completed source files.', 'A successful final stage copies the processed file here and removes it from landing.'),
    ('reject/duplicates/', 'Same-content files under another name or repeat files.', 'The file is routed here; no transaction rows are loaded again.'),
    ('reject/validation_failed/', 'Files that fail structural or data-quality validation.', 'File is retained with failure evidence for investigation and correction.'),
], [2350, 2520, 4950])
doc.add_heading('5.1 Lifecycle decisions', level=2)
for item in [
    'New files must be placed only in landing/. A source file is never overwritten during processing.',
    'A successful pipeline completes Snowflake RAW validation and Gold refresh before archive routing occurs.',
    'A duplicate file name is detected through audit records. A different file name with identical content is detected through a SHA-256 content hash.',
    'A rejected source remains available for troubleshooting; it does not silently disappear.',
]: add_bullet(item)

doc.add_heading('6. Bronze Layer — Raw Ingestion and Auditability', level=1)
doc.add_paragraph('Bronze is the landing representation of the file in Databricks. Its job is to preserve source rows and add operational lineage. It is intentionally not the layer for cleaning or business transformation.')
add_callout('Final raw-data rule', 'Bronze preserves every row from the source file, including exact duplicates. Earlier experimentation used SELECT DISTINCT and reduced the original source count; the final design removes this from Bronze. Exact duplicate removal is performed only in Silver.')
doc.add_heading('6.1 Bronze table', level=2)
add_table(['Object', 'Purpose', 'Examples of stored content'], [
    ('workspace.bronze.transactions_raw', 'Preserves all source transaction columns plus load lineage.', 'Time, V1–V28, Amount, Class, load_ts, source_file, source_s3_uri, batch_id, ingestion_date, pipeline_run_id.'),
    ('workspace.ops.processed_files', 'One operational audit record per Bronze file attempt.', 'Source file/path, ETag, SHA-256, size, status, row counts, timestamps, target table, error message.'),
], [2600, 3100, 4120])
doc.add_heading('6.2 Audit columns and their value', level=2)
add_table(['Audit column', 'Why it exists'], [
    ('load_ts / ingestion_date', 'Shows when the source was ingested and supports load-date reporting.'),
    ('source_file / source_s3_uri', 'Links each transaction to its exact inbound object.'),
    ('batch_id', 'Groups a controlled incremental run from source through Gold.'),
    ('pipeline_run_id', 'Correlates a Databricks processing run with its records.'),
    ('source_etag / content_sha256 / file_size_bytes', 'Identifies the physical S3 object and detects same-content files even when renamed.'),
], [2800, 6320])

doc.add_heading('7. Bronze Validation Notebook', level=1)
doc.add_paragraph('Validation was separated from the Bronze transformation notebook so the transformation remains readable and validation can be rerun or demonstrated independently. The Bronze validation notebook reads the audit table and Bronze table for the selected batch and source file.')
add_table(['Validation', 'Expected result / handling'], [
    ('Source-to-Bronze row reconciliation', 'Source row count equals Bronze row count for the final design. A mismatch is visible as a failed reconciliation, not hidden.'),
    ('Mandatory-field null checks', 'Time, Amount, and Class should be populated for the provided source.'),
    ('Amount validity', 'Negative amount count must be zero for this source.'),
    ('Class validity', 'Class must be 0 or 1 in the historical source.'),
    ('Exact duplicate diagnostic', 'Reports whether duplicate source records exist. It does not delete them in Bronze.'),
    ('Audit status/timestamp check', 'Confirms the file was recorded as SUCCESS with a completion timestamp.'),
], [3120, 6000])

doc.add_heading('8. Silver Layer — Standardisation and Review Features', level=1)
doc.add_paragraph('Silver consumes the selected Bronze batch, performs the quality transformations that make the data suitable for analysis, and carries forward the necessary Bronze audit columns. Unlike Bronze, Silver is the appropriate place to remove exact transaction duplicates.')
doc.add_heading('8.1 Silver transformations', level=2)
add_table(['Transformation', 'Implementation intent', 'Reason'], [
    ('Exact-row deduplication', 'Remove only 100% identical transaction records based on source transaction fields; do not remove valid similar transactions.', 'Prevents duplicate analytical counts while preserving raw evidence in Bronze.'),
    ('Standardised types and names', 'Use consistent SQL data types and readable column naming.', 'Makes downstream SQL, Snowflake, and Streamlit reliable.'),
    ('Synthetic transaction timestamp', 'Add the elapsed Time seconds to a stated project start timestamp.', 'Enables date, hour, time-segment, and daily dashboard analysis.'),
    ('Transaction ID', 'Create a deterministic transaction key.', 'Supports joins, queue display, and controlled incremental checks.'),
    ('Amount bucket and percentile', 'Classify amount bands and rank amount in the batch/population.', 'Supports amount distribution analysis and outlier context.'),
    ('Amount z-score / outlier flag', 'Measure relative amount deviation using a documented statistical rule.', 'Provides a contextual review signal; high amount alone is not proof of fraud.'),
    ('Rapid-repeat flag', 'Identify unusually close transactions in the synthetic time sequence.', 'Provides temporal repetition context for analyst prioritisation.'),
    ('Repeated anonymised-pattern flag', 'Identifies repeat occurrence of a selected V-feature pattern without claiming its business meaning.', 'Adds an explainable technical pattern signal despite anonymised fields.'),
    ('Review score and review priority', 'Combine multiple signals into High, Medium, or Low review priority.', 'Orders analyst work. It is not a fraud prediction or decision.'),
], [2100, 3600, 3420])
doc.add_heading('8.2 Final Silver fields retained for Gold', level=2)
add_table(['Field group', 'Key fields', 'Use in Gold / dashboard'], [
    ('Identity and time', 'transaction_id, transaction_timestamp, transaction_date, transaction_hour, time_of_day, time_band', 'Daily trend, time-window analysis, analyst queue.'),
    ('Amount context', 'amount, amount_bucket, amount_percentile, amount_z_score, amount_outlier_flag', 'Amount analysis and review explanation.'),
    ('Historical label', 'fraud_label', 'Historical reporting only; not used to decide priority.'),
    ('Pattern signals', 'rapid_repeat_flag, repeated_anonymized_pattern_flag', 'Operational triage and review context.'),
    ('Priority output', 'review_score, review_priority', 'Primary sorting/filtering fields for analyst workload.'),
    ('Operational lineage', 'batch_id, source_file, source_s3_uri, ingestion_date, load_ts, pipeline_run_id', 'Trace a downstream record back to the inbound file and run.'),
], [1740, 4300, 3080])
add_callout('Why some exploratory fields were removed', 'Fields such as amount_normalized, decile/quartile/dense-rank, multiple redundant time windows, previous/next value helpers, and unused intermediate pattern keys were removed because they did not drive Gold reporting or explain a review decision. The final model keeps only fields with a demonstrated consumer or operational purpose.', LIGHT_GOLD)

doc.add_heading('9. Silver Validation and Incremental Controls', level=1)
add_table(['Control', 'How it works'], [
    ('Bronze-to-Silver reconciliation', 'Compare selected batch counts; any exact duplicate reduction is explicitly recorded and explainable.'),
    ('Silver audit table', 'workspace.ops.silver_processed_batches records source file, batch, status, Bronze count, Silver count, timestamps, and errors.'),
    ('Batch/source guard', 'The same source batch cannot be transformed into Silver twice as a new load.'),
    ('Feature checks', 'Confirm timestamps, amount fields, review priority, and required columns are present before Snowflake loading.'),
    ('Demo-friendly validation', 'Separate notebook shows before/after counts and selected-batch evidence without cluttering transformation notebooks.'),
], [2800, 6320])

doc.add_heading('10. Snowflake RAW and Gold Analytics', level=1)
doc.add_heading('10.1 Snowflake RAW', level=2)
doc.add_paragraph('FRAUD_DB.RAW.TRANSACTIONS receives only the requested Silver batch. Before writing, the connector checks whether the same batch_id and source_file already exist. If so, it skips the write; if not, it appends. It then reconciles the Snowflake RAW batch count against the Silver batch count.')
doc.add_heading('10.2 Gold outputs', level=2)
add_table(['Gold object', 'Type', 'Contents and value'], [
    ('FRAUD_DAILY_SUMMARY', 'Table', 'Daily transaction count, amount, averages, min/max amounts, historical label counts, priority counts, and rapid-repeat count. Used for daily overview cards/charts.'),
    ('AMOUNT_BUCKET_SUMMARY', 'Table', 'Transaction count/amount and review-priority distribution by amount range. Used for amount analysis.'),
    ('RISK_TIME_DASHBOARD', 'Table', 'Priority distribution by time segment, time window, amount range, and key derived measures. Used for risk-and-time visuals and top-100 validation.'),
    ('ANALYST_REVIEW_QUEUE', 'View', 'Current transaction-level list with transaction time, amount, time labels, flags, score, and review priority. Used for searchable/filterable analyst work.'),
], [2500, 1000, 5620])
doc.add_heading('10.3 Why the analyst queue is a view', level=2)
doc.add_paragraph('The first three Gold objects are materialised aggregate tables because dashboards repeatedly need grouped metrics. The analyst queue remains a view because it should always reflect the latest RAW records and sorting/filtering logic without storing a duplicate copy of every transaction. The Streamlit app can query the current review queue directly.')

doc.add_heading('11. Streamlit Dashboard', level=1)
doc.add_paragraph('The native Snowflake Streamlit app provides the business-consumption layer. It reads only Snowflake Gold outputs, keeps the dashboard separated from transformation logic, and makes review-priority data accessible without direct SQL knowledge.')
add_table(['Dashboard area', 'What it shows', 'Interaction'], [
    ('Summary', 'Total transactions, total transaction amount, High review priority count, and Medium review priority count.', 'KPI cards and priority visuals.'),
    ('Daily Transaction Overview', 'Daily transaction counts and amount movement.', 'Trend visual with readable business labels.'),
    ('Risk and Time Analysis', 'Review priority by time segment/window and day.', 'Priority filters and time-based charting.'),
    ('Amount Analysis', 'Amount-range distribution, totals, and associated priorities.', 'Amount visual and min/max amount range selection.'),
    ('Analyst Review', 'Transaction-level queue ordered by review priority and supporting signals.', 'Priority, time segment, time window, amount range, date, ID search, clear filters, and download.'),
], [2100, 4200, 2820])
add_callout('Dashboard wording', 'Review Priority is intentionally used instead of Fraud or Confirmed Fraud. The dashboard is an operational triage tool; an analyst must review a transaction before any fraud conclusion.')

doc.add_heading('12. Dagster Orchestration', level=1)
doc.add_paragraph('Dagster is mandatory in the final design. It is the control plane that launches the jobs, waits for each dependency, captures execution evidence, and routes the source file based on the final outcome.')
add_code_flow('landing_file_sensor -> run_bronze -> run_silver -> load_snowflake_raw -> validate_snowflake_raw -> refresh_gold -> archive_processed_file')
add_table(['Dagster step', 'Action', 'Success condition'], [
    ('landing_file_sensor', 'Polls S3 landing/ for eligible new CSV files and creates one run request.', 'New content is found, not already audited, and no conflicting active pipeline run is permitted.'),
    ('run_bronze', 'Launches the Bronze Databricks Job with catalog, source path/name, batch ID, content hash, ETag, and size.', 'Bronze raw and audit entry are written.'),
    ('run_silver', 'Launches the Silver Databricks Job for the same batch/source.', 'Silver transaction and Silver audit output are complete.'),
    ('load_snowflake_raw', 'Launches the Silver-to-Snowflake job.', 'New batch appended or safely skipped; count reconciles.'),
    ('validate_snowflake_raw', 'Queries Snowflake for the selected batch.', 'Snowflake RAW rows equal expected Silver rows.'),
    ('refresh_gold', 'Calls ANALYTICS.REFRESH_FRAUD_GOLD().', 'Gold tables and analyst queue are refreshed.'),
    ('archive_processed_file', 'Copies source to archive/ and removes it from landing/.', 'Success lifecycle is visible in S3 and Dagster logs.'),
], [2220, 4050, 2850])
doc.add_heading('12.1 Parameter flow', level=2)
doc.add_paragraph('The sensor creates batch_id from the S3 event context and passes catalog_name, source_file_name, source_path, batch_id, source_etag, content_sha256, and file_size_bytes. These values make every run deterministic and traceable. The earlier manual batch_001 value was a setup/test value; automated runs use dynamically generated S3-based batch IDs.')

doc.add_heading('13. Duplicate Prevention, Idempotency, and Failure Routing', level=1)
add_table(['Scenario', 'Detection method', 'Result'], [
    ('Same name and same source is uploaded again', 'Audit table finds prior successful source S3 URI/name or batch.', 'No reprocessing; duplicate route is used.'),
    ('Different name but identical data content', 'Dagster computes SHA-256 and checks it against processed-file audit records.', 'New filename is sent to reject/duplicates/; no duplicate analytical rows.'),
    ('New file with new content', 'New hash + no successful duplicate record.', 'Normal Bronze -> Silver -> Snowflake -> Gold flow.'),
    ('Input / quality / processing failure', 'Databricks/Dagster exception or validation failure.', 'Run fails visibly; failure evidence is logged and file is routed to reject/validation_failed/ when failure routing is enabled.'),
    ('Dagster retry', 'Retry policy is limited and observable.', 'Transient workload issues can retry; final failure remains explicit for support.'),
], [2360, 3600, 3160])
add_callout('Key defence for review', 'Filename is not treated as proof of uniqueness. The pipeline combines source metadata with a content SHA-256 hash. Therefore an identical file renamed by a user does not create a second Bronze/Silver/Snowflake load.')

doc.add_heading('14. Security and Configuration', level=1)
for item in [
    'AWS IAM follows least privilege: list landing/archive/reject prefixes; read landing files; write archive/reject objects; delete landing objects only after successful archive.',
    'Databricks accesses S3 through an IAM role and external location/storage credential; access is scoped to the required bucket prefixes.',
    'Snowflake password is stored as a Databricks secret for connector execution, not hard-coded in notebooks.',
    'Dagster local .env contains AWS, Databricks, and Snowflake runtime configuration. It is excluded from Git.',
    'Databricks personal-access tokens and AWS access keys are never included in source code, documentation, screenshots, or the repository.',
    'Only read-only, allowlisted Gold objects should be exposed to the optional natural-language-to-SQL website.',
]: add_bullet(item)

doc.add_heading('15. Testing and Validation Evidence', level=1)
add_table(['Test case', 'Evidence to capture for submission/demo'], [
    ('Initial source ingestion', 'S3 landing object; processed_files audit row; Bronze count; Bronze validation output.'),
    ('Silver quality transformation', 'Silver batch audit row; Bronze-to-Silver reconciliation; duplicate diagnostic.'),
    ('Snowflake incremental load', 'RAW count by batch_id and source_file; Snowflake reconciliation result.'),
    ('Gold refresh', 'Procedure result and counts from all three Gold tables plus analyst review view.'),
    ('Archive success path', 'Dagster log showing archive_processed_file and S3 archive/ object.'),
    ('Same-file duplicate protection', 'Same-content/new-name test, duplicate audit evidence, and reject/duplicates/ object.'),
    ('Failure/reject path', 'A controlled invalid-file test, error log, and reject/validation_failed/ object.'),
    ('Top-100 review validation', 'Query the top 100 rows from ANALYST_REVIEW_QUEUE ordered by priority/score and verify the supporting flags.'),
], [2850, 6270])

doc.add_heading('16. Data Model Documentation', level=1)
add_table(['Schema', 'Object', 'Grain', 'Primary purpose'], [
    ('Databricks Bronze', 'transactions_raw', 'One row per source record per accepted source file.', 'Raw preservation and traceability.'),
    ('Databricks Ops', 'processed_files', 'One row per Bronze file attempt.', 'Ingestion audit, file metadata, duplicate prevention.'),
    ('Databricks Silver', 'transactions_silver_incremental', 'One clean, deduplicated row per transaction.', 'Feature-enriched analyst-ready data.'),
    ('Databricks Ops', 'silver_processed_batches', 'One row per Silver batch attempt.', 'Transformation audit and reconciliation.'),
    ('Snowflake RAW', 'TRANSACTIONS', 'One Silver transaction record per loaded batch.', 'Snowflake analytic source.'),
    ('Snowflake Analytics', '3 tables + analyst queue view', 'Aggregate grain varies; queue is transaction-level.', 'Dashboard and analyst consumption.'),
], [1750, 2550, 2560, 4260])

doc.add_heading('17. Repository Structure', level=1)
add_code_flow('fraud-detection-pipeline/\n  databricks/bronze/                 # SQL incremental ingestion notebook\n  databricks/silver/                 # SQL transformation notebook\n  databricks/validations/            # Bronze/Silver validation notebooks\n  snowflake/                          # RAW/Gold/procedure SQL and Streamlit app\n  orchestration/dagster_creditfraud/  # Dagster definitions, sensor, jobs, resources\n  docs/                               # data model, testing evidence, AI usage\n  diagrams/                           # architecture SVG/PNG\n  README.md')

doc.add_heading('18. Final Deliverables Checklist', level=1)
add_table(['Deliverable', 'Status / final action'], [
    ('Working pipeline', 'Core path implemented. Re-run a controlled end-to-end input after final notebook/schema changes.'),
    ('Git repository and README', 'Ensure all notebooks/SQL, diagrams, setup instructions, and .gitignore are committed; exclude secrets.'),
    ('Data model documentation', 'Use Section 16 plus column-level reference in docs/data-model.'),
    ('Test cases and results', 'Capture validations and controlled success/duplicate/failure scenarios.'),
    ('Architecture diagram', 'Use diagrams/final_fraud_detection_architecture_v2.svg and export PNG for presentation.'),
    ('Presentation deck and live demo', 'Use Section 19 runbook and assign each team member a responsibility.'),
    ('AI tool usage write-up', 'State how AI assisted documentation, troubleshooting, design, or code review; keep human validation explicit.'),
    ('Optional bonus website', 'Only begin once core evidence and documentation are complete.'),
], [2900, 6220])

doc.add_heading('19. Recommended Live Demo Flow', level=1)
for item in [
    'Show S3 prefixes: landing, archive, reject/duplicates, and reject/validation_failed.',
    'Show a selected source file and explain the input fields (Time, V1–V28, Amount, Class) and data limitations.',
    'Show Dagster lineage and the ordered pipeline run: Bronze, Silver, Snowflake load/validation, Gold refresh, archive.',
    'Open processed_files and silver_processed_batches to demonstrate batch IDs, source lineage, timestamps, and counts.',
    'Run the separate Bronze and Silver validation notebooks for the selected batch ID.',
    'Show Snowflake RAW count by batch ID, then show Gold object counts and the analyst queue.',
    'Open the Streamlit dashboard; demonstrate KPI cards, filters, range selection, review queue, clear filters, and download.',
    'Show the duplicate-file proof: same content/different file name is rejected by SHA-256 rather than loaded twice.',
    'Conclude with archive evidence and explain that review priority is analyst triage, not a fraud conclusion.',
]: add_number(item)

doc.add_heading('20. Optional Bonus: Natural-Language Query Website', level=1)
doc.add_paragraph('The optional extra-credit extension is a website where a user asks a business question in plain English. The website converts it to safe, read-only SQL and returns results from an approved platform. It must not receive broad account access or arbitrary SQL permissions.')
add_table(['Design rule', 'Requirement'], [
    ('Data scope', 'Allowlist only curated Snowflake Gold tables and ANALYST_REVIEW_QUEUE, or approved Databricks curated data.'),
    ('SQL safety', 'Only SELECT; block DDL/DML, multiple statements, metadata access, and unrestricted table names.'),
    ('Cost/abuse control', 'Limit rows, apply statement timeouts, log prompts and generated SQL, and require a user review step.'),
    ('Secrets', 'Keep credentials on the server side; never ship them to browser code or a public repository.'),
    ('Disclosure', 'Make clear that review-priority answers are not confirmed fraud decisions.'),
], [2500, 6620])

doc.add_heading('21. Current Finalisation Priorities', level=1)
for item in [
    'Keep the final Bronze notebook raw-preserving (no SELECT DISTINCT) and use Silver for exact-duplicate removal only.',
    'Ensure the final Snowflake Gold procedure matches the current final Silver field names, especially review_score and review_priority.',
    'Complete and capture the controlled validation-failure route to reject/validation_failed/.',
    'Run one final fresh, unique test file through the entire orchestration after code changes; capture each layer’s evidence.',
    'Export final architecture PNG, README, data model, test report, deck, and AI-usage note before submission.',
]: add_bullet(item)

doc.add_heading('Appendix A — Terminology', level=1)
add_table(['Term', 'Meaning in this project'], [
    ('Bronze', 'Raw source data plus ingestion metadata; source rows are not altered.'),
    ('Silver', 'Clean and deduplicated transaction data with documented derived features.'),
    ('Gold', 'Business-facing summaries and a current analyst review view in Snowflake.'),
    ('Review priority', 'High/Medium/Low ordering for analyst attention; not a confirmed fraud label.'),
    ('Incremental load', 'Only a new, unprocessed batch/source is added; repeated content is skipped/rejected.'),
    ('Idempotency', 'Repeating a run does not duplicate downstream data.'),
    ('Reconciliation', 'A count comparison between layers to prove the expected records reached the target.'),
], [2400, 6720])

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('End of document')
r.italic = True
r.font.color.rgb = RGBColor.from_string(GRAY)

doc.core_properties.title = 'Credit Card Fraud Detection - End-to-End Data Engineering Project Documentation'
doc.core_properties.subject = 'Architecture, process, operations, validation, and deliverables'
doc.core_properties.author = 'Credit Card Fraud Detection Project Team'
doc.save(OUT)
print(OUT.resolve())
