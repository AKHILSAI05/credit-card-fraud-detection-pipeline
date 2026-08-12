from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).resolve().parent / 'outputs' / 'Fraud_Detection_Pipeline_Change_Log_and_Final_Design.docx'
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = '2E74B5'; DARK = '1F4D78'; NAVY = '0B2545'; LIGHT = 'E8EEF5'; GRAY = 'F2F4F7'; CALLOUT = 'F4F6F9'; GREEN = 'E2F0D9'; GOLD = 'FFF2CC'; RED = 'FCE4D6'; MUTED = '666666'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)

def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

def set_font(run, size=11, color='000000', bold=None, italic=None):
    run.font.name = 'Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def add_text(p, text, size=11, color='000000', bold=None, italic=None):
    r = p.add_run(text); set_font(r, size, color, bold, italic); return r

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.15
    add_text(p, text, 10.5)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.15
    add_text(p, text, 10.5); return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    add_text(p, text, {1:16,2:13,3:12}[level], BLUE if level < 3 else DARK, True)
    return p

def para(doc, text='', bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, 11, '000000', True); add_text(p, text[len(bold_lead):], 11)
    else:
        add_text(p, text, 11)
    return p

def callout(doc, title, text, fill=CALLOUT):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT; set_table_widths(t, [6.5])
    cell = t.cell(0,0); set_cell_shading(cell, fill); set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); add_text(p, title + ' ', 10.5, NAVY, True); add_text(p, text, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def table(doc, headers, rows, widths, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = 'Table Grid'; set_table_widths(t, widths)
    set_repeat_table_header(t.rows[0])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; set_cell_shading(c, LIGHT); set_cell_margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); add_text(p, h, 9.5, NAVY, True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]; set_cell_margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if status_col is not None and i == status_col:
                val = str(value).upper()
                if 'PENDING' in val: set_cell_shading(c, GOLD)
                elif 'VERIFIED' in val or 'IMPLEMENTED' in val: set_cell_shading(c, GREEN)
                elif 'CAUTION' in val: set_cell_shading(c, RED)
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
            add_text(p, str(value), 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return t

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'; styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); styles['Normal'].font.size = Pt(11)
for s, size, color, before, after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
    st=styles[s]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

# Header/footer
header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT; add_text(header, 'Credit Card Fraud Detection | Pipeline Change Log', 8.5, MUTED)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(footer, 'Internal project documentation | Updated 9 August 2026', 8.5, MUTED)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(72); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
add_text(p, 'CREDIT CARD FRAUD DETECTION', 13, DARK, True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
add_text(p, 'Pipeline Change Log and Final Design', 25, NAVY, True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(32)
add_text(p, 'Evolution of the Bronze, Silver, Snowflake Gold, and orchestration design', 13, MUTED)
table(doc, ['Document status', 'As of', 'Scope'], [[ 'Current technical design record', '9 August 2026', 'Bronze, Silver, Gold, Snowflake, Dagster, S3 and dashboard' ]], [1.65,1.1,3.75])
callout(doc, 'Important status:', 'This document records the latest approved design and code direction. Earlier tables still include historical runs from the prior design. A controlled clean rebuild is pending before the final schema and row counts are treated as production-ready.', GOLD)

heading(doc,'1. Executive summary')
para(doc,'The project evolved from a basic one-time ingestion pipeline into an incremental, auditable fraud-risk data platform. The final approach preserves the source exactly in Bronze, standardizes and deduplicates only in Silver, publishes business-friendly analytical objects in Snowflake Gold, and uses Dagster to coordinate, validate, archive, and prevent duplicate work.')
table(doc,['Layer','Final responsibility','Key design decision'],[
    ['S3 landing','Receive source files','Files are processed incrementally, then archived or rejected.'],
    ['Databricks Bronze','Immutable raw ingestion','Preserve every source row; never remove transaction duplicates here.'],
    ['Databricks Silver','Validated, standardized transaction layer','Remove exact duplicate records and calculate transparent review-priority signals.'],
    ['Snowflake RAW','Incremental operational copy of Silver','Append only a new batch after a batch-and-file guard.'],
    ['Snowflake Gold','Dashboard-ready aggregates and analyst queue','Three tables plus one live view, with business-friendly names.'],
    ['Dagster','Orchestration and operational controls','One controlled pipeline at a time; retries, validation, archive and duplicate routing.'],
], [1.15,2.1,3.25])

heading(doc,'2. Current status and important rebuild caveat')
table(doc,['Area','Latest position','Status'],[
 ['Final Bronze SQL','Designed to preserve all source rows and track content hashes.','Implemented in code'],
 ['Final Silver SQL','Lean final schema and exact-duplicate removal design agreed.','Ready for controlled rebuild'],
 ['Snowflake Gold procedure','Business-name Gold object definitions prepared.','Ready to run after RAW reload'],
 ['Dagster pipeline','Bronze → Silver → Snowflake RAW → validation → Gold → archive flow exists.','Verified in earlier runs; latest changes need retest'],
 ['Historic batch_001','Has 283,726 Bronze rows versus 284,807 source rows from an old DISTINCT logic.','Caution: historic data must be rebuilt'],
 ['Clean rebuild','Truncate/reload final objects from source after final notebooks are confirmed.','Pending'],
], [1.55,3.45,1.5], 2)
callout(doc,'Why this matters:', 'The old source load used SELECT DISTINCT in Bronze, removing 1,081 exact duplicate source records. This violated the Bronze principle. The final process fixes this by preserving the raw file entirely in Bronze and removing exact duplicate records only in Silver.', RED)

heading(doc,'3. Bronze layer evolution')
heading(doc,'3.1 Earlier Bronze version',2)
add_bullet(doc,'Loaded the Kaggle creditcard.csv data from S3 landing into workspace.bronze.transactions_raw.')
add_bullet(doc,'Added useful operational metadata such as source file name, S3 path, batch ID, ingestion date, load timestamp, and pipeline run ID.')
add_bullet(doc,'Maintained an audit table in workspace.ops.processed_files.')
add_bullet(doc,'However, SELECT DISTINCT / duplicate removal was performed during Bronze ingestion, reducing the historical original-file count from 284,807 to 283,726. This is the behavior being corrected.')
heading(doc,'3.2 Final Bronze design',2)
table(doc,['Change','Final behavior','Reason'],[
 ['Raw preservation','Use SELECT src.* rather than SELECT DISTINCT.','Bronze is the auditable source-of-truth layer; no genuine source record is discarded.'],
 ['Operational audit fields','load_ts, source_file, source_s3_uri, batch_id, ingestion_date, pipeline_run_id.','Enables lineage, troubleshooting, reconciliation and repeatable batch reporting.'],
 ['Content identity','Record source ETag, file size and SHA-256 content hash in audit records.','Detects same content even if another file name is used.'],
 ['File-level idempotency','Check successful file/content audit records before normal processing.','Avoids running the same source again.'],
 ['Routing','Success moves from landing/ to archive/; repeated content goes to reject/duplicates/.','Landing remains a queue; archive and reject provide operational evidence.'],
 ['Audit outcomes','RUNNING, SUCCESS, FAILED, DUPLICATE_CONTENT with start/end timestamps, counts and errors.','Gives a clear audit trail for reviewers and the team.'],
], [1.35,2.9,2.25])
heading(doc,'3.3 Bronze audit record')
para(doc,'Each file creates or updates one operational record in workspace.ops.processed_files. The audit record contains the source identity, target table, pipeline run, row counts, status, timestamps, error text, and content identity values. Dagster uses this evidence to decide whether a file is new, already processed, or a duplicate under a different name.')

heading(doc,'4. Silver layer evolution')
heading(doc,'4.1 What Silver does in the final design',2)
para(doc,'Silver reads the raw Bronze batch, applies standard data quality checks, removes only exact duplicate transaction records, standardizes fields, creates analysis-ready time and amount attributes, and produces explainable review-priority signals. It does not claim that anonymized V1-V28 columns are customer, card, merchant, or location attributes.')
table(doc,['Field group','Final retained fields / logic','Why retained'],[
 ['Source lineage','source_file, source_s3_uri, batch_id, ingestion_date, load_ts, pipeline_run_id, record_hash','Keeps batch-level traceability through Databricks and Snowflake.'],
 ['Original transaction data','Time, V1-V28 held as feature_01 to feature_28, Amount, Class','Preserves source data; V-features remain anonymized and are not given false business meanings.'],
 ['Identity and time','transaction_id, transaction_timestamp, transaction_date, transaction_hour, time_of_day, time_band','Supports time trends and dashboard filtering; timestamp is synthetically anchored to 1 Aug 2026 plus elapsed Time seconds.'],
 ['Amount analysis','amount, amount_bucket, amount_percentile, amount_z_score, amount_outlier_flag','Supports distribution analysis and explains unusual amount context without claiming that amount alone means fraud.'],
 ['Historical label','fraud_label: Class 0 = Legitimate, Class 1 = Fraud','Used only to validate and describe historic dataset results; never used as the review-priority input.'],
 ['Review signals','rapid_repeat_flag, repeated_anonymized_pattern_flag, review_score, review_priority, review_reason','Creates an explainable analyst queue based on behavior signals rather than Class.'],
], [1.25,3.4,1.85])
heading(doc,'4.2 Removed from the final stored Silver output',2)
para(doc,'The earlier notebook had several fields that were not consumed by Gold, Streamlit, or the review rule. These should be removed from the final persisted Silver table to keep the model understandable and maintainable.')
table(doc,['Removed field / group','Why removed'],[
 ['potential_risk_band; review_required','Duplicated the final review-priority decision without adding a distinct business use.'],
 ['amount_normalized; amount_decile; amount_quartile; amount_dense_rank','Not used by Gold or the dashboard after validation of the final analytical design.'],
 ['transactions_last_5m; transactions_last_15m; amount_last_5m; same_amount_count_last_5m','Not needed for the agreed final scoring model.'],
 ['anonymized_pattern_key; repeated_anonymized_pattern_count','Temporary intermediate logic; retain the boolean flag and reason, not an opaque implementation field.'],
 ['previous_transaction_timestamp; previous_amount; next_transaction_timestamp; next_amount','Intermediate window-function outputs; no need to persist after flags are calculated.'],
 ['Validation helper flags in Silver output','has_nulls, is_valid_amount, is_valid_class, validation_status, validation_reason belong in validation logic/notebooks, not the final analytics table.'],
], [2.65,3.85])
heading(doc,'4.3 Exact-duplicate handling moved to Silver',2)
para(doc,'Bronze retains all records. Silver uses record_hash to identify rows that are exactly identical across the original transaction values. Only one copy of an exact duplicate is retained in Silver. This lets the team reconcile the expected source result: Bronze preserves 284,807 source rows; Silver is expected to contain 283,726 unique transaction records for the original dataset, subject to the final validation run.')
heading(doc,'4.4 Explainable review-priority rule',2)
table(doc,['Signal','Score contribution','Business interpretation'],[
 ['Rapid repeat','+3','A same-amount transaction occurs immediately after the prior globally ordered transaction (within the agreed short time threshold).'],
 ['Repeated anonymized pattern','+2','A repeated pattern made from Amount and selected anonymized features appears in a short window; this is a pattern signal, not a claim about the feature meanings.'],
 ['Very high or very low amount','+2','An amount extreme contributes context, but cannot alone classify an event as fraud.'],
 ['Amount outlier','+1','Amount has an unusual statistical position in the dataset.'],
 ['Contextual rule','+1','Extreme amount at night or rapid repeat during business hours provides additional context.'],
 ['Priority mapping','High >= 5; Medium >= 3; otherwise Low','Prioritizes analyst review; it is not a confirmed fraud prediction.'],
], [1.65,1.25,3.6])

heading(doc,'5. Snowflake RAW and Gold evolution')
heading(doc,'5.1 RAW load controls',2)
add_bullet(doc,'Databricks reads only the selected Silver batch and appends it to FRAUD_DB.RAW.TRANSACTIONS.')
add_bullet(doc,'Before append, it checks whether the same BATCH_ID plus SOURCE_FILE already exists in Snowflake RAW. If it does, the run is safely skipped.')
add_bullet(doc,'After the append, the pipeline reconciles the Silver selected-batch row count with Snowflake RAW for the same batch and source file.')
add_bullet(doc,'The historical backup FRAUD_DB.RAW.TRANSACTIONS_LEGACY_20260807 is preserved and must not be deleted during the clean rebuild.')
heading(doc,'5.2 Gold object naming and design',2)
para(doc,'Earlier Gold definitions used technical names such as SUSPICIOUS_TRANSACTION_RISK. The latest Gold procedure exposes clear business-facing names: REVIEW_PRIORITY, REVIEW_SCORE, and REVIEW_REASON. This makes the Streamlit dashboard, review queue, and presentation easier to explain.')
table(doc,['Gold object','Type','Contents and dashboard use'],[
 ['ANALYTICS.FRAUD_DAILY_SUMMARY','Table','Daily volume, transaction amount, average/min/max amount, historical label counts, priority counts, rapid repeat/repeated pattern/outlier counts.'],
 ['ANALYTICS.AMOUNT_BUCKET_SUMMARY','Table','Amount range-level transaction counts, value totals, historical label count and review-priority signals.'],
 ['ANALYTICS.RISK_TIME_DASHBOARD','Table','Date, time segment/window, amount range and review priority: counts, amounts and behavioral-signal counts.'],
 ['ANALYTICS.ANALYST_REVIEW_QUEUE','View','Live detailed list of transaction-level records, ordered for analyst review with clear aliases.'],
], [2.4,0.75,3.35])
heading(doc,'5.3 Why the analyst queue is a view',2)
para(doc,'The analyst queue is intentionally a view rather than a fourth physical Gold table. It performs no aggregation, should always show the latest current RAW records, and avoids copying every transaction into another stored table. The three Gold tables are materialized because they pre-aggregate dashboard metrics; the view is ideal for a live, filterable analyst work queue.')
heading(doc,'5.4 Gold refresh procedure',2)
para(doc,'ANALYTICS.REFRESH_FRAUD_GOLD() is the single Snowflake procedure called by Dagster after Snowflake RAW validation succeeds. It recreates the three Gold summary tables and the live analyst queue view from RAW. This guarantees that every successful batch refreshes the reporting layer consistently.')

heading(doc,'6. Orchestration, S3 routing and duplicate prevention')
table(doc,['Pipeline step','What happens','Operational control'],[
 ['1. Bronze','Dagster starts the Bronze Databricks Job with catalog, file, path, batch ID and content identity.','Audit log and raw-row preservation.'],
 ['2. Silver','Dagster starts the Silver Job after Bronze succeeds.','Exact duplicate removal only in Silver; Silver audit.'],
 ['3. Snowflake RAW','Dagster starts the Silver-to-Snowflake Job.','Batch + source-file skip and row-count reconciliation.'],
 ['4. RAW validation','Dagster confirms the loaded batch count in Snowflake.','Stops downstream reporting if the count is wrong.'],
 ['5. Gold refresh','Dagster calls ANALYTICS.REFRESH_FRAUD_GOLD().','All dashboard objects refresh together.'],
 ['6. Archive/reject','On success copy to archive then delete landing object; duplicate content goes to reject/duplicates.','Landing behaves as a queue and retains no completed source files.'],
], [1.2,3.2,2.1])
para(doc,'The Dagster sensor scans landing/ approximately every 30 seconds. The final sensor design selects one new candidate per tick and checks for an active run before requesting another, preventing two full pipelines from starting simultaneously. A stored sensor cursor and S3_SENSOR_START_AFTER were also introduced to avoid replaying older files when the sensor is first enabled.')
callout(doc,'Duplicate-file rule:', 'File name alone is not sufficient. The sensor calculates a streaming SHA-256 content hash and compares it with successful audit records. A different file name with identical content is routed to reject/duplicates/ and marked DUPLICATE_CONTENT instead of being ingested.', GREEN)
heading(doc,'6.1 S3 permissions and folder plan',2)
table(doc,['Folder','Purpose','Permission pattern'],[
 ['landing/','New source file queue','List and Get; Delete only after successful archive/reject routing.'],
 ['archive/','Successfully completed source evidence','List and Put.'],
 ['reject/duplicates/','Same content received again under any name','List and Put.'],
 ['reject/validation_failed/','Reserved destination for invalid-source processing failures','List and Put; routing implementation remains pending.'],
], [1.55,3.25,1.7])

heading(doc,'7. Validation and audit improvements')
heading(doc,'7.1 Dedicated validation notebooks',2)
para(doc,'To make the demo and review evidence clearer, validation is being separated from transformation notebooks. A Bronze validation notebook was created with ten cells for batch-specific operational checks. This keeps the Bronze and Silver transformation notebooks focused and makes before/after evidence easy to query during the demo.')
table(doc,['Validation area','Examples of checks'],[
 ['Bronze reconciliation','Source rows vs Bronze rows, audit status, batch/source evidence, load-date and batch counts.'],
 ['Bronze data quality','Null Time/Amount/Class, negative Amount, invalid Class values, exact duplicate group inspection.'],
 ['Silver reconciliation','Bronze batch rows, duplicate removal count, invalid removal count, final Silver count, Silver audit status.'],
 ['Snowflake validation','Silver selected-batch count matches FRAUD_DB.RAW.TRANSACTIONS after guarded append.'],
 ['Gold validation','Top 100 review queue records sorted by priority, review score, amount percentile and transaction time; expected Gold-object counts.'],
], [1.8,4.7])
heading(doc,'7.2 Batch and load-date reporting',2)
para(doc,'Batch_ID-wise and load_date-wise counts were added as evidence requirements. They make it possible to demonstrate exactly which source batch was processed, when it arrived, and how many rows travelled through each layer.')

heading(doc,'8. Streamlit dashboard improvements')
para(doc,'The Snowflake-native Streamlit app was redesigned as a more polished operational dashboard. It uses business-friendly labels rather than underscored database names and makes it clear that review priority is not a confirmed fraud decision.')
table(doc,['Improvement','Latest behavior'],[
 ['KPI cards','Total transactions, total transaction amount, High review priority, and Medium review priority are displayed as cards.'],
 ['Readable labels','Examples: Review Priority, Transaction Time, Amount Range, Time Segment and Time Window.'],
 ['Filters','Review priority, time segment, time window, amount range, transaction date range, transaction-ID search and top-N selection.'],
 ['Amount control','Minimum and maximum amount are included as a range control for clearer analysis.'],
 ['Usability','Clear Filters control and CSV download are available for analyst use.'],
 ['Dashboard sections','Summary, Daily Transaction Overview, Risk & Time Analysis, Amount Analysis and Analyst Review.'],
 ['Analyst queue','A filterable live queue uses the Gold view and clear aliases.'],
], [1.85,4.65])

heading(doc,'9. Controlled final rebuild plan')
para(doc,'Do not remove the Snowflake legacy backup. Once the final Bronze and Silver notebook content is confirmed, perform one controlled clean rebuild so all physical tables match this document.')
for item in [
 'Keep FRAUD_DB.RAW.TRANSACTIONS_LEGACY_20260807 as the historic backup.',
 'Confirm final Bronze has no SELECT DISTINCT and final Silver only persists the agreed required columns.',
 'Truncate or recreate final pipeline tables/audits according to the approved rebuild plan; remove obsolete historical objects only after confirming the backup.',
 'Place the original creditcard.csv in landing/ and run the full Dagster pipeline with the sensor disabled or a controlled manual run.',
 'Validate source 284,807 → Bronze 284,807; validate Silver expected unique count 283,726; reconcile selected batch with Snowflake RAW.',
 'Run ANALYTICS.REFRESH_FRAUD_GOLD(), validate all three Gold tables and the analyst view, then confirm Streamlit.',
 'Test a new valid batch, same-content/different-name duplicate routing, successful archive routing, and the single-active-run sensor guard.',
]: add_number(doc,item)
callout(doc,'Rebuild rule:', 'The final Bronze table is the only layer that must match source row-for-row. The Silver count may be lower only because its explicitly documented exact-duplicate policy is applied there.', GOLD)

heading(doc,'10. Remaining items and verification checklist')
table(doc,['Item','Why it matters','Status'],[
 ['Finalize Silver persisted columns','Ensure only agreed business/analytics fields remain; do not persist intermediate or validation helper columns.','Pending confirmation before rebuild'],
 ['Create Silver validation notebook','Provide clean, repeatable evidence like Bronze validation.','Pending'],
 ['Implement validation-failure routing','Move invalid files to reject/validation_failed/ with a clear audit status.','Pending'],
 ['Retest content-hash duplicate route','Prove different names with same content are rejected without loading.','Pending verification'],
 ['Retest single-active-run guard','Prove sensor does not launch two concurrent pipelines.','Pending verification'],
 ['Run controlled clean rebuild','Replace historical pre-final tables with correct final schema/counts.','Pending'],
 ['Confirm latest Streamlit deployment','Paste/update final app code after Gold procedure executes, then test every tab.','Pending verification'],
 ['Documentation package','Add current notebooks, SQL, Dagster source, validation screenshots and this change log to Git/Notion.','In progress'],
], [2.0,3.05,1.45], 2)

heading(doc,'Appendix A. Key terminology for the final review')
table(doc,['Term','Meaning in this project'],[
 ['Historical fraud label','The original Class field mapped to Fraud/Legitimate. It validates historic outcomes but is not used to create review priority.'],
 ['Review priority','High, Medium or Low operational order for analyst attention based on transparent behavioral/context features.'],
 ['Content hash','SHA-256 fingerprint of a file’s bytes. Same content has the same hash even when the file name changes.'],
 ['Bronze','Raw auditable copy of source data plus lineage metadata.'],
 ['Silver','Validated and standardized transaction layer with exact duplicates removed.'],
 ['Gold','Snowflake analytics tables and live queue designed for BI/dashboard consumption.'],
 ['Idempotent load','A repeated run safely avoids inserting the same logical batch again.'],
], [1.65,4.85])

heading(doc,'Appendix B. Final object reference')
table(doc,['Platform','Object'],[
 ['Databricks Bronze','workspace.bronze.transactions_raw'],
 ['Databricks Bronze audit','workspace.ops.processed_files'],
 ['Databricks Silver','workspace.silver.transactions_silver_incremental'],
 ['Databricks Silver audit','workspace.ops.silver_processed_batches'],
 ['Snowflake RAW','FRAUD_DB.RAW.TRANSACTIONS'],
 ['Snowflake Gold','FRAUD_DB.ANALYTICS.FRAUD_DAILY_SUMMARY'],
 ['Snowflake Gold','FRAUD_DB.ANALYTICS.AMOUNT_BUCKET_SUMMARY'],
 ['Snowflake Gold','FRAUD_DB.ANALYTICS.RISK_TIME_DASHBOARD'],
 ['Snowflake Gold','FRAUD_DB.ANALYTICS.ANALYST_REVIEW_QUEUE'],
 ['Snowflake app','FRAUD_DB.ANALYTICS.FRAUD_RISK_DASHBOARD'],
 ['Orchestration','orchestration/dagster_creditfraud'],
], [2.2,4.3])

doc.core_properties.title = 'Credit Card Fraud Detection - Pipeline Change Log and Final Design'
doc.core_properties.subject = 'Bronze, Silver, Gold, Snowflake and Dagster change record'
doc.core_properties.author = 'Credit Card Fraud Detection Project Team'
doc.save(OUT)
print(OUT)
