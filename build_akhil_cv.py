from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
DOCX_PATH = OUT / "Akhil_Sai_Garikapati_CV.docx"
PDF_PATH = OUT / "Akhil_Sai_Garikapati_CV.pdf"

NAVY = "#202B3B"
NAVY_RGB = RGBColor(32, 43, 59)
TEXT = "#111827"
MUTED = "#4B5563"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def docx_section(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, "202B3B")
    set_cell_margins(cell, top=30, start=30, bottom=30, end=30)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.1)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    return table


def docx_line(doc, parts, size=7.5, after=0.35, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(17, 24, 39)
    return p


def docx_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.13)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(7.1)
    r.font.color.rgb = RGBColor(17, 24, 39)
    return p


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.30)
    section.bottom_margin = Inches(0.40)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(7.4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AKHIL SAI GARIKAPATI")
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = NAVY_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("AI & Data Engineer")
    r.font.name = "Times New Roman"
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = RGBColor(75, 85, 99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("garikapatiakhilsai@gmail.com  |  +91 8008534118  |  Hyderabad, India  |  linkedin.com/in/akhilsai07")
    r.font.name = "Times New Roman"
    r.font.size = Pt(6.1)
    r.font.color.rgb = RGBColor(55, 65, 81)

    docx_section(doc, "PROFILE SUMMARY")
    docx_line(doc, [("Computer Science Engineering graduate specializing in Artificial Intelligence and Intelligent Process Automation, with a 3-month AI & Data Engineering internship at Anblicks. Hands-on exposure to SQL, Python, Snowflake, Databricks, data ingestion, transformation, and cloud workflows.", False, False)], size=7.1, after=1)

    docx_section(doc, "EDUCATION")
    docx_line(doc, [("B.Tech, CSE (AI&IPA)", True, False), ("  |  KL University  |  2026  |  CGPA: 8.78", False, False)], after=0.7)
    docx_line(doc, [("Higher Secondary (XII)", True, False), ("  |  Sri Chaitanya  |  2022  |  Percentage: 83.5%", False, False)], after=3)

    docx_section(doc, "INTERNSHIP EXPERIENCE")
    docx_line(doc, [("Intern - AI & Data Engineer", True, False), ("  |  Anblicks  |  ", False, False), ("May 2026 - Aug 2026", False, True)], after=0.5)
    docx_bullet(doc, "Worked on data engineering projects using SQL, Python, Snowflake, and Databricks.")
    docx_bullet(doc, "Gained hands-on experience in data processing, transformation, ingestion, and cloud-based data workflows.")
    docx_bullet(doc, "Collaborated with industry professionals and gained exposure to real-world data engineering practices.")
    doc.add_paragraph().paragraph_format.space_after = Pt(0.5)

    docx_section(doc, "CAPSTONE PROJECT")
    docx_line(doc, [("Credit Card Fraud Detection Data Engineering Pipeline", True, False)], after=0.5)
    docx_bullet(doc, "Built an automated AWS S3 - Databricks - Snowflake pipeline with Dagster orchestration.")
    docx_bullet(doc, "Implemented incremental loads, audit logs, data validation, duplicate prevention, and archive/reject routing.")
    docx_bullet(doc, "Created Snowflake Gold analytics and a Streamlit dashboard for transaction monitoring and fraud-risk review.")
    doc.add_paragraph().paragraph_format.space_after = Pt(0.5)

    docx_section(doc, "SKILLS")
    docx_line(doc, [("Technical: ", True, False), ("Python, SQL, Databricks, Snowflake, AWS S3, Dagster, Streamlit, Git", False, False)], after=0.7)
    docx_line(doc, [("Soft Skills: ", True, False), ("Problem Solving, Data Analysis, Teamwork, Communication", False, False)], after=3)

    docx_section(doc, "CERTIFICATIONS")
    docx_bullet(doc, "Databricks Certified Data Engineer Associate - Databricks (2026)")
    docx_bullet(doc, "Microsoft Certified Azure AZ-900 - Microsoft (2025)")
    docx_bullet(doc, "Oracle Generative AI Certified - Oracle (2024)")
    doc.add_paragraph().paragraph_format.space_after = Pt(0.5)

    docx_section(doc, "ACHIEVEMENTS & EXTRA-CURRICULARS")
    docx_bullet(doc, "Organized and promoted hackathons and technical events, encouraging student participation and collaboration.")
    docx_bullet(doc, "Participated in the Google Build & Blog Hackathon, collaborating to develop and present a solution.")

    doc.save(DOCX_PATH)


def build_pdf():
    styles = getSampleStyleSheet()
    page_width, _ = letter
    usable = page_width - 2 * 0.45 * inch
    story = []

    name = ParagraphStyle("Name", parent=styles["Normal"], fontName="Times-Bold", fontSize=15, leading=16.5, alignment=TA_CENTER, textColor=colors.HexColor(NAVY), spaceAfter=0)
    role = ParagraphStyle("Role", parent=styles["Normal"], fontName="Times-Italic", fontSize=8, leading=9, alignment=TA_CENTER, textColor=colors.HexColor(MUTED), spaceAfter=0.5)
    contact = ParagraphStyle("Contact", parent=styles["Normal"], fontName="Times-Roman", fontSize=5.8, leading=7, alignment=TA_CENTER, textColor=colors.HexColor(MUTED), spaceAfter=2.5)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName="Times-Roman", fontSize=7.0, leading=8.25, textColor=colors.HexColor(TEXT), spaceAfter=1.0)
    line = ParagraphStyle("Line", parent=body, spaceAfter=0.15)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=14, firstLineIndent=-7, bulletIndent=6, spaceAfter=0.1)
    section_text = ParagraphStyle("SectionText", parent=styles["Normal"], fontName="Times-Bold", fontSize=8.0, leading=8.5, textColor=colors.white)

    def section(title):
        t = Table([[Paragraph(title, section_text)]], colWidths=[usable], rowHeights=[0.14 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(NAVY)),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 0.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5),
        ]))
        story.extend([t, Spacer(1, 0.025 * inch)])

    def bullet_item(text):
        story.append(Paragraph(text, bullet, bulletText="•"))

    story.append(Paragraph("AKHIL SAI GARIKAPATI", name))
    story.append(Paragraph("AI &amp; Data Engineer", role))
    story.append(Paragraph("garikapatiakhilsai@gmail.com &nbsp; | &nbsp; +91 8008534118 &nbsp; | &nbsp; Hyderabad, India &nbsp; | &nbsp; linkedin.com/in/akhilsai07", contact))

    section("PROFILE SUMMARY")
    story.append(Paragraph("Computer Science Engineering graduate specializing in Artificial Intelligence and Intelligent Process Automation, with a 3-month AI &amp; Data Engineering internship at Anblicks. Hands-on exposure to SQL, Python, Snowflake, Databricks, data ingestion, transformation, and cloud workflows.", body))
    story.append(Spacer(1, 0.02 * inch))

    section("EDUCATION")
    story.append(Paragraph("<b>B.Tech, CSE (AI&amp;IPA)</b> &nbsp; | &nbsp; KL University &nbsp; | &nbsp; 2026 &nbsp; | &nbsp; CGPA: 8.78", line))
    story.append(Paragraph("<b>Higher Secondary (XII)</b> &nbsp; | &nbsp; Sri Chaitanya &nbsp; | &nbsp; 2022 &nbsp; | &nbsp; Percentage: 83.5%", body))
    story.append(Spacer(1, 0.02 * inch))

    section("INTERNSHIP EXPERIENCE")
    story.append(Paragraph("<b>Intern - AI &amp; Data Engineer</b> &nbsp; | &nbsp; Anblicks &nbsp; | &nbsp; <i>May 2026 - Aug 2026</i>", line))
    bullet_item("Worked on data engineering projects using SQL, Python, Snowflake, and Databricks.")
    bullet_item("Gained hands-on experience in data processing, transformation, ingestion, and cloud-based data workflows.")
    bullet_item("Collaborated with industry professionals and gained exposure to real-world data engineering practices.")
    story.append(Spacer(1, 0.02 * inch))

    section("CAPSTONE PROJECT")
    story.append(Paragraph("<b>Credit Card Fraud Detection Data Engineering Pipeline</b>", line))
    bullet_item("Built an automated AWS S3 - Databricks - Snowflake pipeline with Dagster orchestration.")
    bullet_item("Implemented incremental loads, audit logs, data validation, duplicate prevention, and archive/reject routing.")
    bullet_item("Created Snowflake Gold analytics and a Streamlit dashboard for transaction monitoring and fraud-risk review.")
    story.append(Spacer(1, 0.02 * inch))

    section("SKILLS")
    story.append(Paragraph("<b>Technical:</b> Python, SQL, Databricks, Snowflake, AWS S3, Dagster, Streamlit, Git", line))
    story.append(Paragraph("<b>Soft Skills:</b> Problem Solving, Data Analysis, Teamwork, Communication", body))
    story.append(Spacer(1, 0.02 * inch))

    section("CERTIFICATIONS")
    bullet_item("Databricks Certified Data Engineer Associate - Databricks (2026)")
    bullet_item("Microsoft Certified Azure AZ-900 - Microsoft (2025)")
    bullet_item("Oracle Generative AI Certified - Oracle (2024)")
    story.append(Spacer(1, 0.02 * inch))

    section("ACHIEVEMENTS &amp; EXTRA-CURRICULARS")
    bullet_item("Organized and promoted hackathons and technical events, encouraging student participation and collaboration.")
    bullet_item("Participated in the Google Build &amp; Blog Hackathon, collaborating to develop and present a solution.")

    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, leftMargin=0.45 * inch, rightMargin=0.45 * inch, topMargin=0.30 * inch, bottomMargin=0.35 * inch)
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
