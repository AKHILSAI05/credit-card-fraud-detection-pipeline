from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('outputs/Credit_Card_Fraud_Detection_Demo_Playbook.docx')
OUT.parent.mkdir(exist_ok=True)
NAVY, BLUE, TEAL, LIGHT, PALE, RED = '0B2545', '2E74B5', '0F766E', 'E8EEF5', 'E6F4F1', '9B1C1C'

doc = Document()
s = doc.sections[0]
s.top_margin = Inches(.72); s.bottom_margin = Inches(.72); s.left_margin = Inches(.8); s.right_margin = Inches(.8)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'; styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
styles['Normal'].font.size = Pt(10.5); styles['Normal'].paragraph_format.space_after = Pt(5); styles['Normal'].paragraph_format.line_spacing = 1.12
for nm, sz, col in [('Title', 28, NAVY), ('Subtitle', 13, '5B6573'), ('Heading 1', 16, BLUE), ('Heading 2', 13, BLUE), ('Heading 3', 11, NAVY)]:
    st = styles[nm]; st.font.name = 'Calibri'; st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); st.font.size = Pt(sz); st.font.color.rgb = RGBColor.from_string(col)
styles['Heading 1'].paragraph_format.space_before = Pt(15); styles['Heading 1'].paragraph_format.space_after = Pt(7)
styles['Heading 2'].paragraph_format.space_before = Pt(10); styles['Heading 2'].paragraph_format.space_after = Pt(5)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet'); p.add_run(text); return p

def para(text='', bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p

def callout(title, body, color=PALE):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0,0); shade(c, color); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = c.paragraphs[0]; r = p.add_run(title + ': '); r.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY); p.add_run(body)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c, LIGHT); c.paragraphs[0].runs[0].bold=True
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    for row in t.rows:
        for c in row.cells:
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph('')
    return t

def heading(text, level=1): doc.add_heading(text, level=level)

# Cover
doc.add_paragraph('CAPSTONE DEMO PLAYBOOK', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Credit Card Fraud Detection\nEnd-to-End Data Engineering Pipeline', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph('Live walkthrough • Presenter script • Mentor Q&A • Pre-demo checks', style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Tools: AWS S3 | Databricks | Snowflake | Dagster | Native Streamlit | SQL | Python', style='Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')
callout('Key message', 'This is a data engineering and fraud-review prioritisation pipeline. It is not an ML fraud prediction model, and High/Medium/Low means review priority rather than confirmed fraud.')
doc.add_page_break()

heading('1. Demo objective and story')
para('Demonstrate how a newly landed transaction file moves safely from AWS S3 to Databricks Bronze and Silver, then to Snowflake RAW and Gold, before becoming visible in the Snowflake-native Streamlit dashboard. Show that the pipeline is automated, auditable, incremental, and protected against duplicate-content files.')
table(['What the mentor should see', 'Evidence to show'], [
['End-to-end automation', 'Dagster sensor, job graph, successful run and event log'],
['Source traceability', 'S3 file, Bronze audit columns and ops.processed_files'],
['Data quality', 'Dedicated Bronze/Silver validation notebooks'],
['Incremental controls', 'Batch ID, content SHA-256, Snowflake guarded append'],
['Business consumption', 'Gold objects and Streamlit analyst review dashboard'],
['Operational controls', 'Archive after success; duplicate/failed file reject routing']
], [2.4,4.6])

heading('2. Recommended live-demo route')
table(['Time', 'Screen / action', 'What to say'], [
['0:00–1:00', 'Architecture diagram', 'We built an end-to-end pipeline from S3 through Databricks and Snowflake to an analyst dashboard.'],
['1:00–2:00', 'AWS S3 prefixes', 'landing holds new files; archive holds successful files; reject holds duplicate or failed files.'],
['2:00–4:00', 'Dagster sensor and successful run', 'Dagster coordinates all systems in dependency order and records every run.'],
['4:00–6:00', 'Databricks Bronze + audit', 'Bronze preserves source truth and adds operational lineage.'],
['6:00–9:00', 'Databricks Silver + validation', 'Silver standardises and enriches data for safe analysis and review prioritisation.'],
['9:00–11:00', 'Snowflake RAW and Gold', 'Only the new batch is appended; Gold creates dashboard-ready summaries and queue.'],
['11:00–13:00', 'Native Streamlit', 'Analysts filter, prioritise, investigate and download review results.'],
['13:00–15:00', 'Duplicate/reject proof and Q&A', 'Same content under a new name is stopped using content hash.']
], [1.0,2.2,3.8])

heading('3. Presenter script')
heading('Opening — about 1 minute', 2)
para('“Good morning. Our project is an end-to-end Credit Card Fraud Detection data engineering pipeline. The objective is not to build an ML model. We built a reliable process that ingests transaction files, preserves raw data, transforms it into analyst-ready data, assigns rule-based review priority, and presents it through Snowflake Streamlit. The main technologies are AWS S3, Databricks, Snowflake, Dagster, SQL, Python, and Streamlit.”')

heading('Architecture — about 2 minutes', 2)
para('“Data starts with the Kaggle ULB Credit Card Fraud dataset. It lands in AWS S3, is ingested to Databricks Bronze, transformed in Silver, loaded to Snowflake RAW, refreshed into Gold analytical objects, and shown through the native Streamlit dashboard. Dagster orchestrates this complete cross-platform flow.”')
callout('Say explicitly', 'The source data has anonymised V1–V28 features and no customer, card, merchant or location identifiers. We therefore do not claim customer-level fraud detection. The system prioritises transactions for analyst review using defensible operational signals.')

heading('S3 landing and file lifecycle — about 2 minutes', 2)
para('“Our S3 bucket is intern-final-project-fraud-detection in ap-south-1. The landing prefix receives new input files. A success moves the file to archive only after Bronze, Silver, Snowflake RAW validation, and Gold refresh finish. Duplicate-content files go to reject/duplicates; failed validation files go to reject/validation_failed.”')
para('“We calculate a SHA-256 content hash. This means a different filename cannot bypass duplicate protection if the file contents are the same.”')

heading('Dagster orchestration — about 3 minutes', 2)
para('“Dagster monitors landing and controls this dependency chain: detect S3 file, run Bronze, run Silver, load Snowflake RAW, validate RAW, refresh Gold, then archive the source file. A downstream step does not begin until its upstream dependency has completed successfully.”')
para('Show the job graph and successful run. “Dagster provides run history, retries, logs, materialisation events, and lineage. Databricks Jobs execute the notebooks; Dagster orchestrates the entire multi-platform workflow.”')

heading('Bronze — about 2 minutes', 2)
para('“Bronze is raw and auditable. It keeps the original Time, V1–V28, Amount, and Class fields and adds load_ts, source_file, source_s3_uri, batch_id, ingestion_date, and pipeline_run_id. We preserve source records in Bronze. Data cleaning and exact-record deduplication are controlled in Silver, not Bronze.”')
para('“The ops.processed_files table records processing status, row counts, timestamps, file ETag, SHA-256 hash, file size, and any error message. This gives end-to-end traceability.”')

heading('Silver — about 3 minutes', 2)
para('“Silver standardises types and creates fields needed for analysis: transaction_id, synthetic transaction timestamp, date and hour, time segment, time window, amount range, amount percentile, amount Z-score, amount outlier indicator, rapid-repeat indicator, repeated anonymised-pattern indicator, review score and review priority.”')
para('“The Class column is renamed to fraud_label and retained only as a historical reporting label. It is never used to create review priority. This avoids data leakage.”')
callout('Best phrasing', 'High, Medium and Low are review priorities, not confirmed fraud decisions. A high amount alone is never treated as fraud; priority combines multiple signals such as unusual amount, rapid repeats, repeated patterns and timing context.')

heading('Snowflake and Gold — about 2 minutes', 2)
para('“Snowflake RAW receives only a new Silver batch. Before append, it checks whether the same batch ID and source file are already present. If present, it safely skips the duplicate; if absent, it appends and reconciles the row count.”')
table(['Gold object', 'Purpose'], [
['FRAUD_DAILY_SUMMARY', 'Daily transaction counts, amounts, review-priority counts and historical-label measures'],
['AMOUNT_BUCKET_SUMMARY', 'Analysis by readable amount range'],
['RISK_TIME_DASHBOARD', 'Review priority by time segment, time window and amount range'],
['ANALYST_REVIEW_QUEUE (view)', 'Live transaction-level queue sorted for analyst investigation']
], [2.6,4.4])
para('“The analyst queue is a view because it should always reflect the latest RAW data without duplicating transaction-level records. The other Gold objects are materialised aggregate tables for fast dashboard use.”')

heading('Streamlit dashboard — about 2 minutes', 2)
para('“The Snowflake-native Streamlit dashboard uses Gold objects only. It provides KPI cards, daily transaction overview, risk and time analysis, amount analysis, and an analyst review queue. Analysts can filter by priority, time segment, time window, date range, amount range, minimum/maximum amount, and transaction ID; they can clear filters and download the filtered results.”')

heading('4. Technical walkthrough')
table(['Layer', 'Main permanent objects', 'What to demonstrate'], [
['S3', 'landing/, archive/, reject/duplicates/, reject/validation_failed/', 'File appears, archive proof and reject policy'],
['Databricks Bronze', 'workspace.bronze.transactions_raw; workspace.ops.processed_files', 'Raw fields + audit columns + file audit'],
['Databricks Silver', 'workspace.silver.transactions_silver_incremental; workspace.ops.silver_processed_batches', 'Standardised/enriched data + review priority'],
['Snowflake RAW', 'FRAUD_DB.RAW.TRANSACTIONS', 'Batch append and count reconciliation'],
['Snowflake Gold', '3 tables + ANALYST_REVIEW_QUEUE view', 'Readable reporting objects and queue'],
['Consumption', 'FRAUD_RISK_DASHBOARD', 'Filters, KPIs, charts and download'],
['Orchestration', 'landing_file_sensor; fraud_pipeline_job', 'Run graph, logs, retries, materialisations']
], [1.35,3.55,2.1])

heading('5. Batch IDs, validation and demo queries')
para('A batch ID represents one processing event. In automated operation Dagster creates a dynamic S3-derived batch ID, such as s3_510c71c5655b. batch_001 was used only during early manual setup and should not be described as the long-term automated approach.')
heading('Find the latest successfully processed batch', 2)
code = "SELECT source_file_name, batch_id, status, started_ts, completed_ts\nFROM workspace.ops.processed_files\nORDER BY started_ts DESC;"
p=doc.add_paragraph(); p.style='Normal'; r=p.add_run(code); r.font.name='Consolas'; r.font.size=Pt(8.5)
heading('What to validate live', 2)
for x in ['Source-to-Bronze row count reconciliation', 'Required field null counts: Time, Amount, Class', 'Negative amounts and invalid Class values', 'Exact-row duplicate diagnostic', 'Bronze-to-Silver reconciliation and Silver audit status', 'Snowflake RAW count for the selected batch', 'Gold objects refreshed and Streamlit data visible']:
    add_bullet(x)

heading('6. Mentor questions and answers')
qa = [
('Are you building a fraud prediction model?', 'No. This is a data engineering capstone with rule-based review prioritisation. The system routes transactions for investigation; it does not declare confirmed fraud.'),
('Why retain Class?', 'Class becomes fraud_label and is used only for historical reporting and validation. It is not an input to review priority, which avoids leakage.'),
('What are V1–V28?', 'They are anonymised PCA-transformed features. We do not invent business meanings such as merchant, customer or location. We retain them for fidelity and limited pattern-based detection.'),
('Does high amount mean fraud?', 'No. Low-value fraud exists. Priority is based on combined signals such as unusual amount, rapid repeat, repeated anonymised pattern and time context.'),
('Why preserve rows in Bronze?', 'Bronze is the source-of-truth audit layer. It must preserve what arrived. Exact-row deduplication is an explicit Silver data-quality transformation.'),
('How do you prevent a duplicate file?', 'Sensor cursor, file audit, S3 ETag, SHA-256 content hash, Snowflake batch/source guard, and archive routing work together.'),
('What if same data arrives under a new filename?', 'SHA-256 is computed from the contents. Same content produces the same hash, so the second file is rejected to reject/duplicates without loading.'),
('Why use Dagster as well as Databricks Jobs?', 'Databricks Jobs run the notebooks. Dagster coordinates S3, Databricks, Snowflake, validation, archive/reject actions, retries and cross-platform lineage.'),
('Why is the analyst queue a view?', 'It must always show current transaction-level results and should not duplicate RAW data. Aggregate Gold objects are tables because dashboards reuse them heavily.'),
('What cannot be modelled with this dataset?', 'There are no customer/card/merchant/location identifiers or real timestamps, so real entity velocity, merchant dimensions and SCDs cannot be claimed.')]
for q,a in qa:
    p=doc.add_paragraph(); p.add_run('Q: '+q).bold=True; p.add_run('\nA: '+a)

heading('7. How to handle common issues during the demo')
table(['Situation', 'Response'], [
['Sensor does not trigger', 'Check Dagster is running, sensor is enabled, AWS credentials are valid, landing prefix is correct, and inspect sensor tick logs.'],
['Databricks notebook says temp view not found', 'Run the notebook from the first cell. bronze_source_raw is a session-specific temporary view created by an earlier cell.'],
['Snowflake load is skipped', 'This is expected if the selected batch and source file already exist. Use a genuinely new batch for append demonstration.'],
['Dashboard is blank after schema change', 'Refresh the Gold procedure after the RAW schema is aligned, then reload Streamlit. Gold and dashboard queries must use the same final column names.'],
['Duplicate file reaches landing', 'The pipeline should route it to reject/duplicates. Explain that this is a successful protection outcome, not a pipeline failure.'],
['A file fails validation', 'Show the error/audit evidence and explain it is routed to reject/validation_failed so unsafe data does not reach analytics.']
], [2.25,4.65])

heading('8. Pre-demo checklist')
for x in [
 'Keep one unique, valid demo CSV ready. Do not use a file that has already been processed.',
 'Keep one renamed duplicate-content file ready only for the duplicate-rejection demonstration.',
 'Start Dagster locally and confirm landing_file_sensor is enabled only when ready.',
 'Confirm AWS access, Databricks serverless/SQL compute, and Snowflake FRAUD_WH warehouse are available.',
 'Confirm the current Snowflake Gold procedure matches the final RAW/Silver column names.',
 'Open architecture, S3, Dagster, Databricks Catalog, Snowflake worksheet, and Streamlit in browser tabs beforehand.',
 'Copy the latest dynamic batch ID from workspace.ops.processed_files before running validations.',
 'Never display AWS keys, Databricks tokens, Snowflake passwords, or .env contents.',
 'Use the phrase review priority, never confirmed fraud, throughout the presentation.'
]: add_bullet(x)

heading('9. Final close — 30 seconds')
para('“To summarise, this project demonstrates a governed, incremental and observable data engineering pipeline. It preserves raw source lineage in Bronze, creates cleaned review-prioritisation data in Silver, serves business-ready Gold analytics in Snowflake, and gives analysts an interactive Streamlit interface. Dagster makes the complete S3-to-dashboard process automated, traceable and safe to rerun.”')
callout('Optional bonus', 'The planned extra-credit website will convert plain-English questions into safe read-only SQL against approved Snowflake Gold objects only. It must allowlist data objects, block DDL/DML and multi-statements, apply row limits, and log prompt-to-SQL activity.', 'FFF7E0')

footer = s.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('Credit Card Fraud Detection Pipeline — Demo Playbook')
doc.save(OUT)
print(OUT.resolve())
